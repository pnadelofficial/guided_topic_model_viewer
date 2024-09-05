import streamlit as st
import pandas as pd
from docx import Document
from datetime import date
from html import escape
import re
from ast import literal_eval
import json
from utils import BM25WithOperators

st.title('Guided Topic Model Viewer')

if 'page_count' not in st.session_state:
    st.session_state['page_count'] = 0

@st.cache_resource
def load_tokenizer():
    return json.load(open("./tokenizer.json"))
tokenizer_dict = load_tokenizer()

st.markdown("""
<style>
    body {
        background-color: #f0f0f0;
    }
    .highlight {
        padding: 2px 0;
        border-radius: 3px;
        transition: all 0.2s;
    }
    .highlight:hover {
        box-shadow: 0 0 10px rgba(0,0,0,0.5);
    }
    /* Custom tooltip */
    .highlight {
        position: relative;
    }
    .highlight::after {
        content: attr(title);
        position: absolute;
        bottom: 100%;
        left: 50%;
        transform: translateX(-50%);
        background-color: #333;
        color: white;
        padding: 5px 10px;
        border-radius: 3px;
        font-size: 14px;
        white-space: nowrap;
        opacity: 0;
        transition: opacity 0.3s;
        pointer-events: none;
    }
    .highlight:hover::after {
        opacity: 1;
    }
</style>
""", unsafe_allow_html=True)

def highlight_text(text, lexical_weights, tokenizer_dict):
    sorted_words = sorted(
        [(tokenizer_dict['model']['vocab'][int(id)][0].replace('▁', ''), weight) 
         for id, weight in lexical_weights.items()],
        key=lambda x: len(x[0]),
        reverse=True
    )

    max_weight = max(weight for _, weight in sorted_words)
    min_weight = min(weight for _, weight in sorted_words)
    weight_range = max_weight - min_weight

    def replace_word(match):
        word = match.group(0)
        for token, weight in sorted_words:
            if word.lower() == token.lower():
                normalized_weight = (weight - min_weight) / weight_range if weight_range > 0 else 1
                hue = int(240 * weight) 
                color = f"hsla({hue}, 90%, 65%, {0.3 + 0.7 * normalized_weight})"
                tooltip = f"Token: {escape(word)}, Lexical weight: {weight:.4f}"
                return f'<span class="highlight" style="background-color: {color};" title="{tooltip}">{escape(word)}</span>'
        return word

    pattern = r'\b(' + '|'.join(re.escape(word) for word, _ in sorted_words) + r')\b'
    highlighted_text = re.sub(pattern, replace_word, text, flags=re.IGNORECASE)

    return highlighted_text

uploaded_file = st.file_uploader("Choose a file")
if uploaded_file is not None:
    df = pd.read_csv(uploaded_file, lineterminator='\n')
    if 'lexical_weights' in df.columns:
        df['lexical_weights'] = df['lexical_weights'].apply(literal_eval)

    if 'topic_df' not in st.session_state:
        st.session_state['topic_df'] = df

highlight_cb = st.checkbox("Highlight text")

if 'topic_df' in st.session_state:
    topic_choice = st.selectbox('Pick a topic to look at more closely', list(st.session_state['topic_df'].topic.unique()))
    # page_size = st.number_input("Results per page", min_value=1, value=5)
    st.markdown(f"#### Topic: **{topic_choice.replace('<|eot_id|>', '')}**")
    selected_df = st.session_state['topic_df'][st.session_state['topic_df'].topic == topic_choice]
    fnames = list(selected_df.filename.unique())
    
    docs = selected_df.chunk.to_list()
    bm25 = BM25WithOperators(docs)
    st.write("Search documents in this topic by keyword")
    query = st.text_input("Enter a query")
    if query:
        results = bm25.search(query)
        for result in results:
            row = selected_df[selected_df.chunk == result['text']]
            lexical_weights = row['lexical_weights'].values[0]
            if highlight_cb:
                highlighted_text = highlight_text(result['text'], lexical_weights, tokenizer_dict)
                st.markdown(f"**Document: {row.filename.values[0]}**")
                st.markdown(highlighted_text.replace("<p>","").replace("</p>",""), unsafe_allow_html=True)
            else:
                st.markdown(f"**Document: {row.filename.values[0]}**")
                st.markdown(result['text'].replace("<p>","").replace("</p>",""), unsafe_allow_html=True)
            st.divider()

    if not query:
        st.write("Or browse by filename")
        tabs = st.tabs([u.replace('.pdf', '').split('/')[-1].replace('"', '') for u in fnames])
        for i, tab in enumerate(tabs):
            with tab:
                # st.session_state['pages'] = [self.limited_results[i:i + to_see] for i in range(0, len(self.limited_results), to_see)]
                for j, row in selected_df[selected_df.filename == fnames[i]].iterrows():
                    text = row['chunk']
                    lexical_weights = row['lexical_weights']
                    if highlight_cb:
                        highlighted_text = highlight_text(text, lexical_weights, tokenizer_dict)
                        st.markdown(highlighted_text.replace("<p>","").replace("</p>",""), unsafe_allow_html=True)
                    else:
                        st.markdown(text.replace("<p>","").replace("</p>",""), unsafe_allow_html=True)
                    st.divider()

    if st.button('Export data for this file or search'):
        with st.spinner("Exporting data"):
            selected_df_for_export = selected_df.drop(columns=['lexical_weights', 'index'])
            if query:
                selected_df_for_export = selected_df_for_export[selected_df_for_export.chunk.isin([r['text'] for r in results])]
            else:
                file_to_export = st.selectbox('Pick a file to export', fnames, format_func=lambda x: x.replace('.pdf', '').split('/')[-1].replace('"', ''))
                selected_df_for_export = selected_df_for_export[selected_df_for_export.filename == file_to_export]
            
            added = query if query else fnames[i]
            xl = selected_df_for_export.to_excel(f"{topic_choice}_{added}_{date.today()}.xlsx", engine='xlsxwriter', index=False)
            with open(f"{topic_choice}_{added}_{date.today()}.xlsx", 'rb') as x_file:
                btn = st.download_button(
                    label="Download data as Excel",
                    data=x_file,
                    file_name=f"{topic_choice}_{date.today()}.xlsx",
                    key='excel_dl'
                )
            
            doc = Document()
            if query:
                doc.add_heading(f"{topic_choice} - search results for '{query}'", 0)
            else:
                doc.add_heading(f"{topic_choice} - {fnames[i]}", 0)
            p = doc.add_paragraph(f"{date.today()}")
            p.italic = True
            for i, row in selected_df_for_export.dropna().iterrows():
                p = doc.add_paragraph(f"Document chunk {i+1}")
                p.bold = True
                doc.add_paragraph(row['filename'])
                text = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+', '', row['chunk'])
                c = doc.add_paragraph(text)
                run = c.add_run()
                run.add_break()
            doc.save(f"{topic_choice}_{added}_{date.today()}.docx")
            with open(f"{topic_choice}_{added}_{date.today()}.docx", 'rb') as w_file:
                btn = st.download_button(
                    label="Download data as Word",
                    data=w_file,
                    file_name=f"{topic_choice}_{added}_{date.today()}.docx",
                    key='word_dl'
                )
        st.divider()
    
    if st.button('Export all data for this topic'):
        with st.spinner("Exporting data"):
            selected_df_for_export = selected_df.drop(columns=['lexical_weights', 'index'])
            xl = selected_df_for_export.to_excel(f"{topic_choice}_{date.today()}.xlsx", engine='xlsxwriter', index=False)
            with open(f"{topic_choice}_{date.today()}.xlsx", 'rb') as x_file:
                btn = st.download_button(
                    label="Download data as Excel",
                    data=x_file,
                    file_name=f"{topic_choice}_{date.today()}.xlsx",
                    key='excel_dl'
                )
        
            doc = Document()
            doc.add_heading(topic_choice, 0)
            p = doc.add_paragraph(f"{date.today()}")
            p.italic = True
            for i, row in selected_df.dropna().iterrows():
                p = doc.add_paragraph(f"Document chunk {i+1}")
                p.bold = True
                doc.add_paragraph(row['filename'])
                text = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+', '', row['chunk'])
                c = doc.add_paragraph(text)
                run = c.add_run()
                run.add_break()
            doc.save(f"{topic_choice}_{date.today()}.docx")
            with open(f"{topic_choice}_{date.today()}.docx", 'rb') as w_file:
                btn = st.download_button(
                    label="Download data as Word",
                    data=w_file,
                    file_name=f"{topic_choice}_{date.today()}.docx",
                    key='word_dl'
                )